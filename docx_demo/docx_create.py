#!/usr/bin/env python3
# -*- coding:utf-8 -*- 
# Created by yaochao at 2018/3/12


from docx import Document
from docx.shared import RGBColor

document = Document()

document.add_heading('【药品名称】', level=1)
document.add_paragraph('通用名称：缬沙坦氨氯地平片(I)')
document.add_paragraph('英文名称:Valsartan and Amlodipine Tablets (I)')
document.add_paragraph().add_run('汉语拼音：Xieshatan Anlvdiping Pian').font.color.rgb = RGBColor(0xFF,0X00,0X00)

document.add_heading('【药品名称】', level=1)
document.add_paragraph('通用名称：缬沙坦氨氯地平片(I)')
document.add_paragraph('英文名称:Valsartan and Amlodipine Tablets (I)')
document.add_paragraph().add_run('汉语拼音：Xieshatan Anlvdiping Pian').font.color.rgb = RGBColor(0xFF,0X00,0X00)

document.add_heading('【药品名称】', level=1)
document.add_paragraph('通用名称：缬沙坦氨氯地平片(I)')
document.add_paragraph('英文名称:Valsartan and Amlodipine Tablets (I)')
document.add_paragraph().add_run('汉语拼音：Xieshatan Anlvdiping Pian').font.color.rgb = RGBColor(0xFF,0X00,0X00)

document.save('create_document.docx')