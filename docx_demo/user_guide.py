#!/usr/bin/env python3
# -*- coding:utf-8 -*- 
# Created by yaochao at 2018/3/8


import docx
import os
import re
from execjs_demo.execjs_demo import get_diff
from pprint import pprint
from docx import Document
from docx.shared import RGBColor

base = '/Users/yaochao/work/说明书差异对比和提取'
file1 = os.path.join(base, 'doc_handle1/【JYY编号1】硝苯地平控释片说明书 OCR版-质控后.docx')
file2 = os.path.join(base, 'doc_handle2/【JYY编号1】1.钙拮抗剂-1：[拜新同]硝苯地平控释片（done） 3.docx')
file461 = os.path.join(base, 'doc_handle1/【JYY编号46】氯沙坦钾氢氯噻嗪片说明书 OCR版--质控后.docx')
file462 = os.path.join(base, 'doc_handle2/【JYY编号46】氯沙坦钾氢氯噻嗪片.md')

# 所有文件title的映射表
titles_map = [['核准日期'],
              ['修改日期'],
              ['警告'],
              ['药品名称', '药品名称:'],
              ['商品名'],
              ['通用名'],
              ['成份', '成 份', '主要成份', '主要成分'],
              ['适应症', '适应症/功能主治'],
              ['性状', '性 状'],
              ['规格', '规 格', '规格型号'],
              ['用法用量', '用法和用量'],
              ['按照患者体重50kg计算', '按照患者体重50 kg计算'],
              ['不良反应', '不良反应:'],
              ['禁忌', '禁 忌'],
              ['注意事项', '注意事 项', '注意事项:'],
              ['服药与进食'],
              ['孕妇及哺乳期妇女用药', '孕妇与哺乳期妇女用药', '孕妇及晡乳期妇女用药', '孕妇和哺乳期妇女用药', '孕妇和哺乳期妇女药', '孕妇及哺乳期服妇女用药'],
              ['妊娠期药物危险性等级'],
              ['哺乳期药物危险性等级'],
              ['儿童用药', '儿童用量'],
              ['老年用药', '老年患者用药'],
              ['药物相互作用'],
              ['药物过量'],
              ['临床试验'],
              ['药理毒理', '药物毒理', '药理作用', '毒理研究'],
              ['药代动力学', '药物（代谢）动力学'],
              ['药物分类'],
              ['贮藏', '贮 藏', '贮   藏', '贮裁'],
              ['包装', '包 装'],
              ['有效期', '有 效 期'],
              ['执行标准'],
              ['批准文号', '批准文号:'],
              ['生产企业', '制造厂'],
              ['进口分装企业'],
              ['进口药品注册证号'],
              ['热线']]


def get_all_texts(dir_path):
    files_name = os.listdir(dir_path)
    files = [os.path.join(dir_path, x) for x in files_name]
    texts = []
    for f in files:
        if '.docx' in f:
            doc = docx.Document(f)
            paragraphs = doc.paragraphs
            text = '\n'.join([x.text for x in paragraphs])
            texts.append(text)
        if '.md' in f:
            text = ''
            parts = get_all_parts_md(f)
            for i in parts:
                title = '【' + i[0] + '】'
                text += title
                text += i[1]
            texts.append(text)
    return texts


def get_text(f):
    doc = docx.Document(f)
    paragraphs = doc.paragraphs
    text = '\n'.join([x.text for x in paragraphs])
    return text


def get_all_titles(texts):
    title_re = '【(.+?)】'
    title_re = re.compile(title_re)
    titles = []
    for t in texts:
        r = re.findall(title_re, t)
        titles.extend(r)
    return set(titles)


def get_titles():
    '''
    获取两个文件夹里面的所有docx文档的titles
    :return: 打印出来的是两个文件夹里面的所有docx文档的titles不在titles_map里面的。
    '''
    texts = get_all_texts(dir_path='/Users/yaochao/work/说明书差异对比和提取/doc_handle1')
    titles1 = get_all_titles(texts)

    texts2 = get_all_texts(dir_path='/Users/yaochao/work/说明书差异对比和提取/doc_handle2')
    titles2 = get_all_titles(texts2)
    all_titles = titles1.union(titles2)
    all_titles1 = []
    for i in titles_map:
        all_titles1 += i
    for i in all_titles:
        if i not in all_titles1:
            print(i)


def get_all_parts(f):
    '''
    得到所有的title以及对应的content，(title, content)称为一个part
    :param f:
    :return:
    '''
    if '.docx' in f:
        return get_all_parts_docx(f)
    elif '.md' in f:
        return get_all_parts_md(f)
    else:
        return ''


def get_all_parts_docx(f):
    text = get_text(f)
    title_re = '(?<![和见照])(?<!和\s)(?<!见\s)(?<!照\s)【(.+?)】([\w\W]+?)(?<![和见照])(?<!和\s)(?<!见\s)(?<!照\s)(?=【|$)'
    parts = re.findall(title_re, text)
    return parts


def get_all_parts_md(f):
    title_re = '## (.+?)\n([\w\W]+?)(?=##|$)'
    with open(f, encoding='utf-8') as f:
        text = f.read()
        parts = re.findall(title_re, text)
        return parts


def title_unification(title):
    '''
    把给定的title归一为统一的title
    '''
    for i in titles_map:
        if title in i:
            return i[0]
    return title


def compare_two_files(f1, f2):
    # 1. 分别拆成不同的部分，归一标题
    parts1 = get_all_parts(f1)
    parts1 = [[title_unification(x[0]), x[1]] for x in parts1]
    parts1_titles = [x[0] for x in parts1]
    parts2 = get_all_parts(f2)
    parts2 = [[title_unification(x[0]), x[1]] for x in parts2]
    parts2_titles = [x[0] for x in parts2]
    # 2. for循环，对两个文本做合并，1为f1单独有，2为f2单独有，0为共有
    for idx, i in enumerate(parts1_titles):
        if i in parts2_titles:
            parts1_i = parts1[idx]
            parts1.pop(idx)
            # 把f1和f2中共有的title所对应的part，从f2中移除，方面循环parts2_titles时，剩下的全部是f2中特有的。
            parts2_i_index = parts2_titles.index(i)
            parts2_i = parts2[parts2_i_index]
            parts2_titles.pop(parts2_i_index)
            parts2.pop(parts2_i_index)
            # 组成新的parts1_i，即parts1中第i个元素，第一个是0，代表共有，第二个是parts1中的，第三个是parts2中的，重组后的形式 [0, '生产企业', '\n企业名称：Bayer Pharma AG \n生产地址：51368 Leverkusen,Germany \n电话号码：+49(0)214 30/57430 \n传真号码：+49(9)214 9657430\n', '拜耳医药保健有限公司\n']
            parts1_i = [0, parts1_i[0], parts1_i[1], parts2_i[1]]
            parts1.insert(idx, parts1_i)
        else:
            parts1[idx].insert(0, 1)
    for idx, i in enumerate(parts2_titles):
        parts2[idx].insert(0, 2)
        parts1.insert(0, parts2[idx])

    # 3. for 循环，对比文本的不同
    for i in parts1:
        if i[0] == 0:
            title = i[1]
            text1 = i[2].strip()
            text2 = i[3].strip()
            diff = get_diff(text1, text2)
            i.append(diff)
    # TODO 对parts1排序，按照要求的title的顺序
    new_parts1 = []
    print('parts1:', len(parts1))
    for i in titles_map:
        for ii in parts1:
            if ii[1] == i[0]:
                new_parts1.append(ii)
                parts1.remove(ii)
    print('new_parts1:', len(new_parts1))

    for i in parts1:
        if i not in new_parts1:
            print('new_parts1中缺少的:', i[0], i[1])
    return new_parts1


def get_RGBColor(diff_type):
    '''
    根据diff_type获取RGBColor对象
    '''
    color_map = [RGBColor(0x00, 0X00, 0X00), RGBColor(0xFF, 0X00, 0X00), RGBColor(0x00, 0XFF, 0X00)]
    return color_map[int(diff_type)]


def out_docx(parts):
    '''
    把比较之后的结果输出至docx文档
    :param parts: 比较之后的结果是个数组
    :return:
    '''
    document = Document()
    # 第一种情况，当两个文档有共同title时，即i[0]==0时
    for i in parts:
        if i[0] == 0:
            title = i[1]
            document.add_heading('【' + title + '】', level=1)
            diff = i[-1]
            for ii in diff:
                diff_type = ii[0]
                diff_text = ii[1]
                document.add_paragraph().add_run(diff_text).font.color.rgb = get_RGBColor(diff_type)
    document.save('out.docx')


def get_HTMLColorText(type, text):
    colortext = ['<font color=\'black\'>' + text + '</font>', '<font color=\'red\'>' + text + '</font>',
                 '<font color=\'green\'>' + text + '</font>']
    return colortext[int(type)]


def out_html(parts, out):
    '''
    把比较之后的结果输出至docx文档
    :param parts: 比较之后的结果是个数组
    :return:
    '''
    # 第一种情况，当两个文档有共同title时，即i[0]==0时，第二种情况i[0]==1，第三种i[0]==2。
    html = ''
    for i in parts:
        if i[0] == 0:
            title = i[1]
            html += '<h2>【' + title + '】</h2>'
            diff = i[-1]
            for ii in diff:
                diff_type = ii[0]
                diff_text = ii[1]
                html += get_HTMLColorText(diff_type, diff_text)
        elif i[0] == 1:
            title = i[1]
            html += '<h2><font color="green">【' + title + '】</font></h2>'
            content = i[2]
            html += '<font color="green">' + content + '</font>'
        elif i[0] == 2:
            title = i[1]
            html += '<h2><font color="red">【' + title + '】</font></h2>'
            content = i[2]
            html += '<font color="red">' + content + '</font>'

    with open(out, 'w', encoding='utf-8') as f:
        f.write(html.replace('\n', '<br>'))


def main():
    parts = compare_two_files(file461, file462)
    out_html(parts, out='out46.html')


def go():
    dir_path1 = '/Users/yaochao/work/说明书差异对比和提取/doc_handle1'
    dir_path2 = '/Users/yaochao/work/说明书差异对比和提取/doc_handle2'
    files_name1 = os.listdir(dir_path1)
    files_name1 = [os.path.join(dir_path1, x) for x in files_name1]

    files_name2 = os.listdir(dir_path2)
    files_name2 = [os.path.join(dir_path2, x) for x in files_name2]
    new_files_name1 = []
    for i in files_name1:
        if '.docx' in i or '.md' in i:
            i_re = '/(【JYY编号\d{1,3}】)\w+'
            pre_i = re.findall(i_re, i)
            pre_i = pre_i[0] if pre_i else ''
            new_files_name1.append([pre_i, i])
    new_files_name2 = []
    for i in files_name2:
        if '.docx' in i or '.md' in i:
            i_re = '/(【JYY编号\d{1,3}】)\w+'
            pre_i = re.findall(i_re, i)
            pre_i = pre_i[0] if pre_i else ''
            new_files_name2.append([pre_i, i])
    compare_files = []
    for i in new_files_name1:
        for ii in new_files_name2:
            if i[0] == ii[0]:
                compare_files.append([i[0], i[1], ii[1]])
    assert len(new_files_name1) == len(new_files_name2) == len(compare_files)
    for i in compare_files:
        f1 = i[1]
        f2 = i[2]
        out = i[0]
        print('正在比对：', i[0])
        parts = compare_two_files(f1, f2)
        out_html(parts, out='/Users/yaochao/work/说明书差异对比和提取/out/' + out + '.html')


if __name__ == '__main__':
    go()
